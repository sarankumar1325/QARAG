"""Document processing service - handles PDFs, DOCX, Markdown, text files, and URLs with lazy initialization"""

import io
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Tuple, Optional
from urllib.parse import urlparse
import requests
from bs4 import BeautifulSoup
import pdfplumber
from docx import Document as DocxDocument

from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import get_settings
from app.models import DocumentType, DocumentStatus


class DocumentChunk:
    """Represents a chunk of a document"""

    def __init__(self, content: str, metadata: dict):
        self.content = content
        self.metadata = metadata
        self.id = str(uuid.uuid4())


class DocumentProcessor:
    """Service for processing various document types with lazy initialization"""

    def __init__(self):
        self._settings = None
        self._text_splitter = None
        self._documents_dir = None

    @property
    def settings(self):
        """Lazy load settings"""
        if self._settings is None:
            self._settings = get_settings()
        return self._settings

    @property
    def text_splitter(self):
        """Lazy load text splitter"""
        if self._text_splitter is None:
            self._text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.settings.chunk_size,
                chunk_overlap=self.settings.chunk_overlap,
                separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
            )
        return self._text_splitter

    @property
    def documents_dir(self):
        """Lazy load documents directory"""
        if self._documents_dir is None:
            self._documents_dir = Path("documents")
            self._documents_dir.mkdir(exist_ok=True)
        return self._documents_dir

    def process_file(
        self, file_content: bytes, filename: str, doc_id: str
    ) -> Tuple[DocumentStatus, List[DocumentChunk], Optional[str]]:
        """
        Process an uploaded file and return chunks

        Returns:
            Tuple of (status, chunks, error_message)
        """
        try:
            # Determine document type
            doc_type = self.get_document_type(filename)

            # Extract text based on document type
            if doc_type == DocumentType.PDF:
                text = self._extract_pdf(file_content)
            elif doc_type == DocumentType.DOCX:
                text = self._extract_docx(file_content)
            elif doc_type == DocumentType.MARKDOWN:
                text = self._extract_markdown(file_content)
            elif doc_type == DocumentType.TEXT:
                text = self._extract_text(file_content)
            else:
                return DocumentStatus.FAILED, [], f"Unsupported file type: {filename}"

            if not text or not text.strip():
                return DocumentStatus.FAILED, [], "No text content extracted from file"

            # Clean and chunk text
            cleaned_text = self._clean_text(text)
            chunks = self._chunk_text(cleaned_text, doc_id, filename, doc_type.value)

            # Save file locally
            self._save_file(file_content, filename, doc_id)

            return DocumentStatus.COMPLETED, chunks, None

        except Exception as e:
            return DocumentStatus.FAILED, [], str(e)

    def process_url(
        self, url: str, doc_id: str
    ) -> Tuple[DocumentStatus, List[DocumentChunk], Optional[str]]:
        """
        Process a website URL and return chunks

        Returns:
            Tuple of (status, chunks, error_message)
        """
        try:
            # Validate URL
            if not self._is_valid_url(url):
                return DocumentStatus.FAILED, [], "Invalid URL format"

            # Fetch and parse
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()

            # Extract text from HTML
            text = self._extract_html(response.content, url)

            if not text or not text.strip():
                return DocumentStatus.FAILED, [], "No text content extracted from URL"

            # Clean and chunk text
            cleaned_text = self._clean_text(text)
            chunks = self._chunk_text(cleaned_text, doc_id, url, "html")

            return DocumentStatus.COMPLETED, chunks, None

        except requests.RequestException as e:
            return DocumentStatus.FAILED, [], f"Failed to fetch URL: {str(e)}"
        except Exception as e:
            return DocumentStatus.FAILED, [], str(e)

    def get_document_type(self, filename: str) -> DocumentType:
        """Determine document type from filename"""
        ext = Path(filename).suffix.lower()

        type_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".txt": DocumentType.TEXT,
            ".text": DocumentType.TEXT,
        }

        return type_map.get(ext, DocumentType.TEXT)

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF"""
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(io.BytesIO(content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    def _extract_markdown(self, content: bytes) -> str:
        """Extract text from Markdown"""
        return content.decode("utf-8", errors="ignore")

    def _extract_text(self, content: bytes) -> str:
        """Extract text from plain text file"""
        # Try different encodings
        encodings = ["utf-8", "utf-16", "ascii", "latin-1"]
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore")

    def _extract_html(self, content: bytes, url: str) -> str:
        """Extract text from HTML"""
        soup = BeautifulSoup(content, "lxml")

        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        # Get text
        text = soup.get_text(separator="\n")

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)

        return text

    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove control characters
        text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]", "", text)

        # Fix common PDF extraction issues
        text = re.sub(r"\n\s*\n", "\n\n", text)

        return text.strip()

    def _chunk_text(
        self, text: str, doc_id: str, source: str, doc_type: str
    ) -> List[DocumentChunk]:
        """Split text into chunks"""
        chunks = self.text_splitter.split_text(text)

        document_chunks = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "doc_id": doc_id,
                "source": source,
                "doc_type": doc_type,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "processed_at": datetime.utcnow().isoformat(),
            }
            document_chunks.append(DocumentChunk(chunk, metadata))

        return document_chunks

    def _is_valid_url(self, url: str) -> bool:
        """Validate URL format"""
        try:
            result = urlparse(url)
            return all([result.scheme, result.netloc])
        except Exception:
            return False

    def _save_file(self, content: bytes, filename: str, doc_id: str):
        """Save uploaded file locally"""
        doc_dir = self.documents_dir / doc_id
        doc_dir.mkdir(exist_ok=True)

        filepath = doc_dir / filename
        with open(filepath, "wb") as f:
            f.write(content)


# Create instance - but it won't initialize until actually used
document_processor = DocumentProcessor()
